#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档分块处理模块

基于原始word_chunking_and_neo4j.py中的WordChunker类实现
提供Word文档的智能分块功能
"""

import docx
import os
import re
from typing import List


class WordChunker:
    """用于将Word文档分块处理的类"""

    def __init__(self, max_tokens=2000, overlap_tokens=200):
        """
        初始化分块器

        Args:
            max_tokens: 每个分块的最大token数
            overlap_tokens: 相邻分块之间的重叠token数
        """
        self.max_tokens = max_tokens
        self.overlap_tokens = overlap_tokens

    def estimate_tokens(self, text: str) -> int:
        """
        估算文本中的token数量（简单估算，每个单词/标点符号算一个token）

        Args:
            text: 输入文本

        Returns:
            估算的token数量
        """
        # 简单估算：按照中文每个字符1个token，英文每个单词1个token
        words = re.findall(r'\b\w+\b', text)
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        punctuations = re.findall(r'[^\w\s]', text)

        return len(words) + len(chinese_chars) + len(punctuations)

    def chunk_document(self, doc_path: str) -> List[str]:
        """
        将Word文档分块

        Args:
            doc_path: Word文档路径

        Returns:
            分块后的文本列表
        """
        # 打开Word文档
        doc = docx.Document(doc_path)

        # 获取文档全文
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"

        # 分块
        chunks = []
        current_chunk = ""
        current_token_count = 0

        # 按段落处理
        for para in doc.paragraphs:
            para_text = para.text + "\n"
            para_tokens = self.estimate_tokens(para_text)

            # 如果当前段落加上已有内容会超过最大token数，则结束当前分块
            if current_token_count + para_tokens > self.max_tokens and current_chunk:
                chunks.append(current_chunk)
                # 保留重叠部分作为新分块的开始
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + para_text
                current_token_count = self.estimate_tokens(current_chunk)
            else:
                current_chunk += para_text
                current_token_count += para_tokens

        # 添加最后一个分块
        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def _get_overlap_text(self, text: str) -> str:
        """
        从文本末尾获取重叠部分，优化处理中英文混合文本。

        Args:
            text: 输入文本

        Returns:
            重叠部分文本
        """
        # 优先尝试获取最后一个完整句子或段落作为重叠 (以常见标点或换行符结尾)
        # 匹配到最后一个标点符号/换行符及其之后的所有内容
        paragraph_match = re.search(r'([。？！.!?\n][^。？！.!?\n]*)$', text.rstrip())
        if paragraph_match:
            paragraph_overlap = paragraph_match.group(1).lstrip()
            # 如果段落重叠部分token数不超过最大重叠token数，则使用段落重叠
            if self.estimate_tokens(paragraph_overlap) <= self.overlap_tokens:
                # print(f"\nUsing paragraph overlap: {paragraph_overlap[:50]}...") # Debug
                return paragraph_overlap

        # 如果段落太长或未找到，则按token数从后往前取
        overlap_text = ""
        current_tokens = 0
        # 从后往前遍历字符，构建后缀
        for i in range(len(text) - 1, -1, -1):
            char = text[i]
            # 构造逐步增加的后缀字符串
            potential_overlap = char + overlap_text
            # 估算这个后缀的token数
            estimated_tokens = self.estimate_tokens(potential_overlap)

            # 如果添加当前字符后的token数不超过目标重叠数，则接受它
            if estimated_tokens <= self.overlap_tokens:
                overlap_text = potential_overlap
                current_tokens = estimated_tokens
            else:
                # 如果添加当前字符就超了，说明之前的overlap_text是能满足条件的最长后缀
                break

        # print(f"\nUsing token overlap: {overlap_text[:50]}...") # Debug
        # 移除可能由从后向前构建过程引入的前导空白
        return overlap_text.lstrip()

    def extract_text_from_document(self, doc_path: str) -> str:
        """
        从Word文档中提取纯文本内容
        
        Args:
            doc_path: Word文档路径
            
        Returns:
            提取的文本内容
        """
        doc = docx.Document(doc_path)
        full_text = ""
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空段落
                full_text += para.text + "\n"
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text += " | ".join(row_text) + "\n"
        
        return full_text.strip()

    def chunk_document_with_tables(self, doc_path: str) -> List[str]:
        """
        将Word文档分块，包含表格处理
        
        Args:
            doc_path: Word文档路径
            
        Returns:
            分块后的文本列表
        """
        doc = docx.Document(doc_path)
        chunks = []
        current_chunk = ""
        current_token_count = 0
        
        # 获取文档中所有元素（段落和表格）的顺序
        elements = []
        
        # 遍历文档的所有元素
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                # 找到对应的段落对象
                for para in doc.paragraphs:
                    if para._element == element:
                        elements.append(('paragraph', para.text))
                        break
            elif element.tag.endswith('tbl'):  # 表格
                # 找到对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        # 提取表格文本
                        table_text = ""
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                table_text += " | ".join(row_text) + "\n"
                        elements.append(('table', table_text))
                        break
        
        # 按元素顺序处理分块
        for element_type, element_text in elements:
            if not element_text.strip():
                continue
                
            element_text_with_newline = element_text + "\n"
            element_tokens = self.estimate_tokens(element_text_with_newline)
            
            # 如果当前元素加上已有内容会超过最大token数，则结束当前分块
            if current_token_count + element_tokens > self.max_tokens and current_chunk:
                chunks.append(current_chunk.strip())
                # 保留重叠部分作为新分块的开始
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + element_text_with_newline
                current_token_count = self.estimate_tokens(current_chunk)
            else:
                current_chunk += element_text_with_newline
                current_token_count += element_tokens
        
        # 添加最后一个分块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks


# 便捷函数
def chunk_word_document(doc_path: str, max_tokens: int = 2000, overlap_tokens: int = 200, include_tables: bool = True) -> List[str]:
    """
    便捷函数：将Word文档分块
    
    Args:
        doc_path: Word文档路径
        max_tokens: 每个分块的最大token数
        overlap_tokens: 相邻分块之间的重叠token数
        include_tables: 是否包含表格处理
        
    Returns:
        分块后的文本列表
    """
    chunker = WordChunker(max_tokens=max_tokens, overlap_tokens=overlap_tokens)
    
    if include_tables:
        return chunker.chunk_document_with_tables(doc_path)
    else:
        return chunker.chunk_document(doc_path)


def extract_text_from_word(doc_path: str) -> str:
    """
    便捷函数：从Word文档中提取文本
    
    Args:
        doc_path: Word文档路径
        
    Returns:
        提取的文本内容
    """
    chunker = WordChunker()
    return chunker.extract_text_from_document(doc_path)